"""
TalentScan ATS — Intelligent Resume Analyzer
Powered by Google Gemini API (FREE — 1500 req/day)
"""

import streamlit as st
import base64
import json
import re
import logging
import time
from pathlib import Path
from io import BytesIO

try:
    from openai import OpenAI
    GROQ_OK = True
except ImportError:
    GROQ_OK = False

try:
    import fitz
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False

try:
    import docx
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from PIL import Image
    PIL_OK = True
except ImportError:
    PIL_OK = False

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# ══════════════════════════════════════
#  CONFIG — Groq (GRATUIT, 14400 req/jour, pas de CB)
#  Clé sur : https://console.groq.com  → API Keys → Create
# ══════════════════════════════════════
GROQ_API_KEY = "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"
GROQ_BASE_URL = "https://api.groq.com/openai/v1"
MODEL_NAME   = "meta-llama/llama-4-scout-17b-16e-instruct"  # Vision + gratuit
MAX_PDF_PAGES = 8


@st.cache_resource
def get_model():
    if not GROQ_OK:
        st.error("Installez: pip install openai")
        st.stop()
    if "VOTRE_CLE" in GROQ_API_KEY:
        st.error("⚠️ Ajoutez votre clé dans app.py ligne_API_KEY")
        st.stop()
    return OpenAI(api_key=GROQ_API_KEY, base_url=GROQ_BASE_URL)


# ══════════════════════════════════════
#  FILE PROCESSING
# ══════════════════════════════════════

def pdf_to_pil(pdf_bytes):
    if not PYMUPDF_OK or not PIL_OK:
        return []
    imgs = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for i in range(min(len(doc), MAX_PDF_PAGES)):
            pix = doc[i].get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            imgs.append(Image.frombytes("RGB", [pix.width, pix.height], pix.samples))
        doc.close()
    except Exception as e:
        logger.error(f"PDF→PIL: {e}")
    return imgs


def pdf_text(pdf_bytes):
    if not PYMUPDF_OK:
        return ""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        t = "\n".join(p.get_text("text") for p in doc).strip()
        doc.close()
        return t
    except:
        return ""


def docx_text(b):
    if not DOCX_OK:
        return ""
    try:
        d = docx.Document(BytesIO(b))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except:
        return ""


def prepare(uploaded_file):
    raw = uploaded_file.read()
    ext = Path(uploaded_file.name).suffix.lower()
    text, imgs, mode = "", [], "text"
    if ext == ".pdf":
        text = pdf_text(raw)
        if len(text.split()) < 50:
            imgs = pdf_to_pil(raw)
            mode = "vision"
    elif ext in (".docx", ".doc"):
        text = docx_text(raw)
    elif ext in (".png", ".jpg", ".jpeg", ".webp"):
        imgs = [Image.open(BytesIO(raw))] if PIL_OK else []
        mode = "vision"
    return text, imgs, mode


# ══════════════════════════════════════
#  PROMPT
# ══════════════════════════════════════

SCHEMA = """{
  "candidate": {
    "name": "Prénom Nom",
    "email": "email ou Non trouvé",
    "phone": "tel ou Non trouvé",
    "location": "Ville, Pays ou Non spécifié",
    "linkedin": "URL ou Non trouvé",
    "github": "URL ou Non trouvé",
    "years_experience": "X ans",
    "education": "Diplôme, École",
    "current_role": "Poste actuel",
    "languages": ["Français", "Anglais"]
  },
  "score": 74,
  "score_breakdown": {
    "technical_skills": 80,
    "experience": 70,
    "education": 75,
    "soft_skills": 65,
    "job_fit": 78
  },
  "verdict": "Profil solide",
  "decision": "À interviewer",
  "summary": "Résumé de 2-3 phrases.",
  "matched_skills": ["Python", "Docker"],
  "missing_skills": ["Kubernetes"],
  "extra_skills": ["Rust"],
  "experience_highlights": ["3 ans chez XYZ"],
  "red_flags": ["Trou d emploi non expliqué"],
  "recommendations": [
    {"type": "good", "text": "Point fort"},
    {"type": "tip",  "text": "Conseil"},
    {"type": "warn", "text": "Alerte"}
  ]
}"""


def build_prompt(job_desc, text="", vision=False):
    extra = f"\n\nTEXTE DU CV:\n{text[:5000]}" if text else ""
    vis = "\n\n[Images du CV jointes — lis tout le contenu visible y compris les zones graphiques]" if vision else ""
    return f"""Tu es TalentScan, un ATS expert. Réponds UNIQUEMENT en JSON valide, sans backticks, sans texte avant ou après.

POSTE:
{job_desc}
{extra}{vis}

Retourne exactement ce JSON:
{SCHEMA}

Commence par {{ et termine par }}. Aucun autre texte."""


# ══════════════════════════════════════
#  API CALLS
# ══════════════════════════════════════

def call_text(model, prompt):
    r = model.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=2048,
    )
    return r.choices[0].message.content


def call_vision(model, prompt, imgs):
    import base64, io
    img_messages = []
    for img in imgs[:3]:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        img_messages.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}})
    img_messages.append({"type": "text", "text": prompt})
    r = model.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": img_messages}],
        temperature=0.1,
        max_tokens=2048,
    )
    return r.choices[0].message.content


def parse_json(raw):
    cleaned = re.sub(r"```(?:json)?", "", raw).strip().strip("`").strip()
    try:
        return json.loads(cleaned)
    except:
        m = re.search(r'\{[\s\S]+\}', cleaned)
        if m:
            try:
                return json.loads(m.group())
            except:
                pass
    raise ValueError("L'IA n'a pas retourné un JSON valide. Réessayez.")


def analyze(uploaded_file, job_desc):
    model = get_model()
    text, imgs, mode = prepare(uploaded_file)
    raw = ""
    if mode == "vision" and imgs:
        try:
            raw = call_vision(model, build_prompt(job_desc, text, True), imgs)
        except Exception as e:
            logger.warning(f"Vision failed: {e}, fallback text")
            if text:
                raw = call_text(model, build_prompt(job_desc, text, False))
            else:
                raise
    else:
        raw = call_text(model, build_prompt(job_desc, text, False))
    return parse_json(raw)


# ══════════════════════════════════════
#  UI HELPERS
# ══════════════════════════════════════

def sc(s): return "#1db954" if s>=70 else "#f59e0b" if s>=45 else "#ef4444"
def sl(s): return "🟢 Excellent" if s>=80 else "🟢 Bon candidat" if s>=65 else "🟡 Potentiel" if s>=50 else "🟠 Insuffisant" if s>=35 else "🔴 Non prioritaire"
def de(d):
    d = (d or "").lower()
    return "📞" if "interview" in d else "⭐" if "retenu" in d else "❌" if "non" in d else "🤔"

def pills(skills, color, bg):
    if not skills:
        st.caption("Aucune")
        return
    st.markdown(" ".join(f'<span style="background:{bg};color:{color};border:1px solid {color}40;padding:5px 13px;border-radius:100px;font-size:13px;margin:3px;display:inline-block;">{s}</span>' for s in skills), unsafe_allow_html=True)

def gauge(score):
    c = sc(score); r=54; circ=339.3; off=circ*(1-score/100)
    st.markdown(f"""<div style="text-align:center;padding:16px 0 8px;">
      <svg width="140" height="140" viewBox="0 0 140 140">
        <circle cx="70" cy="70" r="{r}" fill="none" stroke="#f0f0f0" stroke-width="10"/>
        <circle cx="70" cy="70" r="{r}" fill="none" stroke="{c}" stroke-width="10"
          stroke-dasharray="{circ:.1f}" stroke-dashoffset="{off:.1f}"
          stroke-linecap="round" transform="rotate(-90 70 70)"/>
        <text x="70" y="64" text-anchor="middle" font-size="28" font-weight="800" fill="{c}">{score}%</text>
        <text x="70" y="82" text-anchor="middle" font-size="12" fill="#999">match</text>
      </svg>
      <div style="font-size:15px;font-weight:600;color:{c};margin-top:4px;">{sl(score)}</div>
    </div>""", unsafe_allow_html=True)

def bar(label, val):
    c = sc(val)
    st.markdown(f"""<div style="margin-bottom:12px;">
      <div style="display:flex;justify-content:space-between;font-size:13px;margin-bottom:5px;">
        <span style="color:#555;">{label}</span><span style="font-weight:700;color:{c};">{val}%</span>
      </div>
      <div style="background:#f0f0f0;border-radius:100px;height:7px;">
        <div style="width:{val}%;background:{c};border-radius:100px;height:7px;"></div>
      </div>
    </div>""", unsafe_allow_html=True)

def css():
    st.markdown("""<style>
    @import url('https://fonts.googleapis.com/css2?family=Syne:wght@700;800&family=DM+Sans:wght@300;400;500&display=swap');
    #MainMenu,footer,header{visibility:hidden}
    .block-container{padding-top:1.5rem;max-width:1080px}
    div.stButton>button[kind="primary"]{background:#0a0a0f!important;color:white!important;border:none!important;border-radius:100px!important;font-weight:600!important;padding:14px 32px!important;font-size:15px!important;width:100%}
    div.stButton>button[kind="primary"]:hover{background:#e8642a!important}
    div.stButton>button:not([kind="primary"]){border-radius:100px!important;border:1.5px solid #ddd!important}
    div.stButton>button:not([kind="primary"]):hover{border-color:#0a0a0f!important;background:#0a0a0f!important;color:white!important}
    .stProgress>div>div{background:#e8642a!important;border-radius:100px!important}
    .stFileUploader>div{border-radius:16px!important;border:2px dashed rgba(232,100,42,0.35)!important}
    .stTextArea>div>textarea{border-radius:14px!important;font-size:14px!important;border:1.5px solid #e5e5e5!important}
    </style>""", unsafe_allow_html=True)


# ══════════════════════════════════════
#  PAGES
# ══════════════════════════════════════

def page_landing():
    st.markdown("""
    <div style="background:linear-gradient(135deg,#0a0a0f 0%,#16213e 60%,#0a0a0f 100%);
        border-radius:24px;padding:72px 48px;text-align:center;margin-bottom:28px;position:relative;overflow:hidden;">
      <div style="position:absolute;top:-80px;right:-80px;width:350px;height:350px;background:radial-gradient(circle,rgba(232,100,42,0.18),transparent);border-radius:50%;"></div>
      <div style="display:inline-block;background:rgba(232,100,42,0.12);border:1px solid rgba(232,100,42,0.3);color:#e8642a;padding:6px 18px;border-radius:100px;font-size:11px;letter-spacing:2px;text-transform:uppercase;margin-bottom:28px;">
        ✦ ATS Intelligent — AI
      </div>
      <div style="font-family:'Syne',sans-serif;font-size:clamp(40px,6vw,70px);font-weight:800;line-height:0.95;letter-spacing:-2px;color:#f5f2eb;margin-bottom:24px;">
        Recrutez<br><span style="color:#e8642a;">plus vite.</span><br><span style="opacity:0.35;">Plus précis.</span>
      </div>
      <div style="font-size:17px;color:rgba(245,242,235,0.5);font-weight:300;line-height:1.75;max-width:520px;margin:0 auto 12px;">
        Analyse CVs Canva, PDFs scannés et images grâce à la Vision IA — 100% gratuit.
      </div>
    </div>""", unsafe_allow_html=True)

    cols = st.columns(4)
    for col, (n, l) in zip(cols, [("99%","Précision OCR"),("<15s","Analyse"),("Gratuit","0€ / mois"),("Vision","IA native")]):
        col.markdown(f'<div style="background:#0a0a0f;color:#f5f2eb;border-radius:16px;padding:22px;text-align:center;"><div style="font-family:Syne,sans-serif;font-size:28px;font-weight:800;color:#e8642a;">{n}</div><div style="font-size:11px;color:#8a8578;text-transform:uppercase;letter-spacing:1px;margin-top:4px;">{l}</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    feats = [("👁️","Vision OCR IA","Lit les CVs Canva, scans et PDFs imagés grâce à Gemini Vision."),
             ("🧠","Analyse sémantique","Comprend le contexte réel, pas juste des mots-clés."),
             ("🎯","Score pondéré","5 critères : compétences, expérience, formation, soft skills, adéquation."),
             ("⚡","Hybride & Fallback","Texte + Vision combinés. Bascule automatiquement si besoin."),
             ("📊","Rapport complet","Forces, lacunes, red flags, recommandations et décision finale.")]
    c = st.columns(3)
    for i, (ic, t, d) in enumerate(feats):
        c[i%3].markdown(f'<div style="background:white;border:1.5px solid #f0f0f0;border-radius:16px;padding:24px;margin-bottom:16px;min-height:130px;"><div style="font-size:28px;margin-bottom:10px;">{ic}</div><div style="font-family:Syne,sans-serif;font-weight:700;font-size:15px;margin-bottom:8px;">{t}</div><div style="font-size:13px;color:#777;line-height:1.6;">{d}</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    _, c2, _ = st.columns([1,2,1])
    with c2:
        if st.button("🚀 Commencer l'analyse →", type="primary"):
            st.session_state.page = "analyzer"
            st.rerun()
    st.markdown('<div style="text-align:center;margin-top:20px;font-size:12px;color:#bbb;">Propulsé par <strong></strong></div>', unsafe_allow_html=True)


def page_analyzer():
    cb, ct = st.columns([1,6])
    with cb:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("← Retour"):
            st.session_state.page = "landing"
            st.session_state.pop("result", None)
            st.rerun()
    with ct:
        st.markdown('<div style="font-family:Syne,sans-serif;font-size:32px;font-weight:800;letter-spacing:-1px;margin-top:8px;">Analyser un CV</div><div style="color:#888;font-size:14px;margin-bottom:8px;">Uploadez le CV et collez la description du poste — l\'IA fait le reste.</div>', unsafe_allow_html=True)

    st.markdown("---")

    col_cv, col_job = st.columns(2, gap="large")
    with col_cv:
        st.markdown("**📄 CV du candidat**")
        st.caption("PDF (Canva inclus), DOCX, PNG, JPG, WEBP")
        uf = st.file_uploader("CV", type=["pdf","docx","doc","png","jpg","jpeg","webp"], label_visibility="collapsed")
        if uf:
            st.success(f"✅ **{uf.name}** — {len(uf.getvalue())//1024} KB")
            if uf.type.startswith("image/") and PIL_OK:
                st.image(uf, use_container_width=True)
            elif uf.type == "application/pdf" and PYMUPDF_OK and PIL_OK:
                try:
                    rb = uf.read(); uf.seek(0)
                    doc = fitz.open(stream=rb, filetype="pdf")
                    pix = doc[0].get_pixmap(matrix=fitz.Matrix(1.2,1.2))
                    img = Image.frombytes("RGB",[pix.width,pix.height],pix.samples)
                    st.image(img, caption=f"Page 1/{len(doc)}", use_container_width=True)
                    doc.close()
                except: pass

    with col_job:
        st.markdown("**💼 Description du poste**")
        st.caption("Collez l'offre d'emploi ou décrivez le profil recherché")
        jd = st.text_area("JD", height=280, placeholder="Ex: Développeur Full-Stack Senior — React, Node.js, Docker, 3+ ans d'expérience, Agile, anglais courant...", label_visibility="collapsed", key="jd")
        if jd:
            wc = len(jd.split())
            st.caption(f"{wc} mots — {'🟢 Bonne' if wc>40 else '🟡 Courte' if wc>15 else '🔴 Trop courte'}")

    st.markdown("---")

    _, c2, _ = st.columns([1,2,1])
    with c2:
        clicked = st.button("🔍 Analyser avec l'IA", type="primary")

    if clicked:
        errs = []
        if not uf: errs.append("Veuillez uploader un CV.")
        if not jd or len(jd.strip()) < 20: errs.append("Description du poste trop courte.")
        if errs:
            for e in errs: st.error(f"⚠️ {e}")
        else:
            prog = st.progress(0, text="🔄 Initialisation...")
            try:
                if not GROQ_OK: raise ImportError("openai manquant — pip install openai")
                prog.progress(20, text="📖 Lecture du fichier...")
                time.sleep(0.2)
                prog.progress(50, text="🧠 Analyse Gemini Vision en cours...")
                uf.seek(0)
                result = analyze(uf, jd)
                prog.progress(100, text="✅ Terminé !")
                time.sleep(0.3)
                st.session_state.result = result
                prog.empty()
            except ImportError:
                prog.empty()
                st.error("❌ Installez: pip install openai")
            except ValueError as e:
                prog.empty()
                st.error(f"❌ {e}")
            except Exception as e:
                prog.empty()
                es = str(e)
                if "API_KEY" in es or "api key" in es.lower(): st.error("❌ Clé API invalide. Vérifiez API_KEY")
                elif "quota" in es.lower() or "429" in es: st.error("❌ Quota Gemini atteint. Attendez quelques minutes.")
                elif "400" in es: st.error("❌ Fichier non supporté ou trop grand. Essayez un autre format.")
                else: st.error(f"❌ Erreur: {es}")
                with st.expander("Détails"): st.code(es)

    if "result" in st.session_state:
        r = st.session_state.result
        score = max(0, min(100, int(r.get("score", 0))))
        cand = r.get("candidate", {})

        st.markdown("---")
        st.markdown("## 📊 Résultats")

        cg, cb2 = st.columns([1,2], gap="large")
        with cg: gauge(score)
        with cb2:
            st.markdown(f"""<div style="background:#0a0a0f;border-radius:20px;padding:28px 32px;color:#f5f2eb;">
              <div style="font-family:'Syne',sans-serif;font-size:22px;font-weight:800;margin-bottom:6px;">{cand.get('name','Candidat')}</div>
              <div style="font-size:13px;color:#8a8578;margin-bottom:14px;">{cand.get('current_role','')}</div>
              <div style="background:rgba(232,100,42,0.15);border:1px solid rgba(232,100,42,0.3);color:#e8642a;display:inline-block;padding:5px 16px;border-radius:100px;font-size:12px;font-weight:600;margin-bottom:16px;">
                {de(r.get('decision',''))} {r.get('decision','')} — {r.get('verdict','')}
              </div>
              <div style="font-size:14px;color:rgba(245,242,235,0.6);line-height:1.75;">{r.get('summary','')}</div>
            </div>""", unsafe_allow_html=True)

        st.markdown("---")

        bd = r.get("score_breakdown", {})
        if bd:
            st.markdown("### 📈 Détail du score")
            lm = {"technical_skills":"🔧 Compétences","experience":"💼 Expérience","education":"🎓 Formation","soft_skills":"🤝 Soft Skills","job_fit":"🎯 Adéquation"}
            c1, c2 = st.columns(2)
            for i, (k,v) in enumerate(bd.items()):
                with (c1 if i%2==0 else c2): bar(lm.get(k,k), int(v))
            st.markdown("---")

        st.markdown("### 🛠️ Compétences")
        cm, cx, ce = st.columns(3)
        with cm: st.markdown("**✅ Matchées**"); pills(r.get("matched_skills",[]),"#15803d","rgba(29,185,84,0.08)")
        with cx: st.markdown("**❌ Manquantes**"); pills(r.get("missing_skills",[]),"#b91c1c","rgba(239,68,68,0.06)")
        with ce: st.markdown("**➕ Bonus**"); pills(r.get("extra_skills",[]),"#1d4ed8","rgba(42,110,232,0.07)")

        st.markdown("---")
        ci, ch = st.columns(2, gap="large")
        with ci:
            st.markdown("### 👤 Informations candidat")
            for key, val in [("📧 Email",cand.get("email")),("📱 Tél",cand.get("phone")),("📍 Lieu",cand.get("location")),("⏱️ Exp.",cand.get("years_experience")),("🎓 Formation",cand.get("education")),("🔗 LinkedIn",cand.get("linkedin")),("💻 GitHub",cand.get("github")),("🌍 Langues",", ".join(cand.get("languages",[])) or None)]:
                if val and val not in ("Non trouvé","Non spécifié",""):
                    st.markdown(f'<div style="display:flex;justify-content:space-between;padding:9px 0;border-bottom:1px solid #f5f5f5;font-size:14px;"><span style="color:#888;font-size:12px;text-transform:uppercase;">{key}</span><span style="font-weight:500;text-align:right;max-width:60%;word-break:break-word;">{val}</span></div>', unsafe_allow_html=True)
        with ch:
            st.markdown("### 🏆 Points forts & Alertes")
            for h in r.get("experience_highlights",[]): st.markdown(f'<div style="background:#f0fdf4;border:1px solid #bbf7d0;border-radius:10px;padding:10px 14px;margin-bottom:8px;font-size:14px;">✨ {h}</div>', unsafe_allow_html=True)
            for f in r.get("red_flags",[]): st.markdown(f'<div style="background:#fef2f2;border:1px solid #fecaca;border-radius:10px;padding:10px 14px;margin-bottom:8px;font-size:14px;">🚩 {f}</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 💡 Recommandations")
        styles = {"good":("✅","#f0fdf4","#bbf7d0"),"tip":("💡","#fffbeb","#fde68a"),"warn":("⚠️","#fef2f2","#fecaca")}
        for rec in r.get("recommendations",[]):
            ic, bg, bd2 = styles.get(rec.get("type","tip"),styles["tip"])
            st.markdown(f'<div style="background:{bg};border:1px solid {bd2};border-radius:12px;padding:13px 18px;margin-bottom:10px;font-size:14px;line-height:1.65;">{ic} {rec.get("text","")}</div>', unsafe_allow_html=True)

        st.markdown("---")
        cd, cr = st.columns(2)
        with cd:
            st.download_button("📥 Exporter (JSON)", data=json.dumps({"candidat":cand,"score":score,"verdict":r.get("verdict"),"decision":r.get("decision"),"score_breakdown":r.get("score_breakdown"),"competences_matchees":r.get("matched_skills"),"competences_manquantes":r.get("missing_skills"),"recommandations":r.get("recommendations")}, ensure_ascii=False, indent=2), file_name=f"rapport_{cand.get('name','candidat').replace(' ','_')}.json", mime="application/json", use_container_width=True)
        with cr:
            if st.button("🔄 Analyser un autre CV", use_container_width=True):
                st.session_state.pop("result", None)
                st.rerun()


# ══════════════════════════════════════
#  MAIN
# ══════════════════════════════════════
def main():
    st.set_page_config(page_title="TalentScan — ATS Intelligent", page_icon="🎯", layout="wide", initial_sidebar_state="collapsed")
    css()
    if "page" not in st.session_state:
        st.session_state.page = "landing"
    if st.session_state.page == "landing":
        page_landing()
    else:
        page_analyzer()

if __name__ == "__main__":
    main()
